'''
@Time : 2023-10-07 15:43
@Author : laolao
@FileName: Flow cleaning.py
'''

def cleanse_poc(string):
    import re
    # 定义一个正则表达式，用于匹配ip地址的格式
    white_list=[
     r"\bAccept\s*:\s*\S+",
    r"\bCLIENT_IP\s*:\s*\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}",
    r"\bHost\s*:\s*\S+",
    r"\bVia\s*:\s*\S+",
    r"\bAccept-Encoding\s*:\s*.*",
    r"\bAccept-Language\s*:\s*\S+",
    r"\bAccept\s*:\s*.*",
    r"\bReferer\s*:\s*\S+",
    r"\bContent-Length\s*:\s*\d+",
    r"\bContent-Type\s*:\s*\S+",
    r"\bUser-Agent\s*:\s*.*",
    r"\bConnection\s*:\s*.*",
    r"\bContent-Type\s*:\s*\S+",
    r"X-Forwarded-For: ((25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)\.){3}(25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)$",
    r"Host: \d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}:\d{1,5}",
    r"\bHost\s*:\s*\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}",
    r"\bHTTP/1\.[01]",
    r"\b\w+://[^ ]+",
     r"(\s+\S\s+|\s+)"]
    for i in white_list:
        string = re.sub(i, '', string)
    return ''.join(string.split('\n'))

# %%cmd
# pip install python-docx

import os
current_path ='G:\\实习\\张家口重保\\'
import docx

my_code=[]
for file in os.listdir(current_path):
    file_path = os.path.join(current_path, file)
    if os.path.isfile(file_path):
        doc = docx.Document(file_path)
        poc = doc.paragraphs[9].text+"\n"+doc.paragraphs[10].text
        my_code.append(cleanse_poc(poc))
with open("data/nsfocus_isop.txt", "w",encoding='gbk') as f:
    f.write('\n'.join(my_code))